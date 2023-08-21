import openai
import os
from dotenv import load_dotenv, find_dotenv
from docx import Document
import streamlit as st
from PIL import Image

# Load environment variables from API_key.env and retrieve the API key
load_dotenv(find_dotenv())
openai.api_key = os.getenv('OPENAI_API_KEY')

# Function to get a chat completion response from the model
def get_completion(prompt, model="gpt-3.5-turbo", temperature=0):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=temperature
    )
    return response.choices[0].message["content"]

# Function to translate text to the target language
def translate_text(text, target_language):
    response = openai.Completion.create(
        engine="davinci",  # Use the Davinci engine for translation
        prompt=f"Translate the following text to {target_language}:\n{text}",
        max_tokens=100
    )
    translated_text = response.choices[0].text.strip()
    return translated_text

# Function to replace names based on the selected country
def replace_names(text, country):
    name_mapping = {
        "India": {"Alex": "Rahul", "Smith": "Sharma"},
        "German": {"Alex": "Andre", "Smith": "Müller"},
        "Spanish": {"Alex": "Javier", "Smith": "González"},
        # Add more country mappings as needed
    }
    for original_name, localized_name in name_mapping.get(country, {}).items():
        text = text.replace(original_name, localized_name)
    return text

# Function to localize and translate the demo script
def localize_and_translate_doc(doc, target_language, country):
    doc = Document(doc)
    script_text = "\n".join([para.text for para in doc.paragraphs])
    
    localized_text = replace_names(script_text, country)
    translated_text = translate_text(localized_text, target_language)
    
    localized_doc = Document()
    for para in translated_text.split("\n"):
        localized_doc.add_paragraph(para)
    
    localized_doc.save("localized_demo_script.docx")

# Streamlit UI
st.title("Demo Script Localization")
uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

if uploaded_file:
    target_language = st.selectbox("Select Target Language", ["French", "German", "Spanish"])  # Add more options
    country = st.selectbox("Select Your Country", ["India", "Germany", "Spain"])  # Add more options

    if st.button("Localize and Download"):
        # Perform localization and save the localized document
        localize_and_translate_doc(uploaded_file, target_language, country)
        st.success("Document localized and downloaded!")
